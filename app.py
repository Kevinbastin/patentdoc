import streamlit as st
import os
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------ LOCAL MODULE IMPORTS (Phi-3) ---------------
from generate_title import generate_title_from_abstract
from generate_claims import generate_claims_from_abstract
from generate_summary import summarize_abstract
from generate_field_of_invention import generate_field_of_invention
from generate_background import generate_background_locally
from generate_detailed_description import generate_detailed_description
from generate_brief_description import generate_brief_description
from generate_summary_of_drawings import generate_drawing_descriptions as generate_drawing
from generate_objects import generate_objects_of_invention
from cpc_classifier import classify_cpc
from export_to_pdf import create_patent_pdf

# -------------- UI: HEADER & DISCLAIMER ------------------
st.title("🧠 PatentDoc Co-Pilot")

with st.expander("⚠️ Legal Disclaimer & AI Usage Notice", expanded=False):
    st.markdown("""
    This tool uses local AI language models (Phi-3, Llama.cpp, etc.) to help draft patent documents.

    ⚠️ **Warning:** The AI-generated content may contain factual errors, legal inaccuracies, or formatting that does not comply with USPTO guidelines.
    Always consult a qualified **patent attorney** before relying on the generated material for filing or legal use.

    📜 **Model License Notice:** The underlying LLMs (e.g., Phi-3) are subject to their respective open-source licenses.

    🔍 **Responsibility Disclaimer:** This app is a research prototype. It is **not** a substitute for professional legal services or patent counsel.
    """)

# ---------------- MAIN INPUT FIELDS ---------------------
abstract = st.text_area("📄 Enter Invention Abstract", height=200)
st.session_state["abstract_input"] = abstract
drawing_summary = st.text_area("🎨 Enter Drawing Summary (optional)", height=150)

# ----------------- SESSION STATE INIT -------------------
for key in [
    "title", "claims", "summary", "field_of_invention",
    "background","objects_of_invention", "detailed_description", "brief_description", "summary_drawings", "cpc_result"
]:
    if key not in st.session_state:
        st.session_state[key] = ""

# ------------------- GENERATION BUTTONS -----------------
# ------------------- GENERATION BUTTONS (FIXED) -----------------
# ------------------- GENERATION BUTTONS (FIXED) -----------------
if st.button("📌 Generate Title"):
    with st.spinner("Generating title..."):
        try:
            result = generate_title_from_abstract(abstract)
            if isinstance(result, dict):
                st.session_state.title = result.get("title", "")
            else:
                st.session_state.title = result
            st.success("Done!")
        except Exception as e:
            st.error(f"❌ Title generation failed: {e}")

if st.session_state.title:
    with st.expander("📘 Title"):
        st.write(st.session_state.title)


if st.button("🔖 Generate Claims"):
    with st.spinner("Generating claims..."):
        try:
            result = generate_claims_from_abstract(abstract)
            if isinstance(result, dict):
                st.session_state.claims = result.get("text", result.get("claims", ""))
            else:
                st.session_state.claims = result
            st.success("Done!")
        except Exception as e:
            st.error(f"❌ Claim generation failed: {e}")

if st.session_state.claims:
    with st.expander("🧾 Claims"):
        st.write(st.session_state.claims)


if st.button("🧷 Generate Summary"):
    with st.spinner("Generating summary..."):
        try:
            result = summarize_abstract(abstract)
            if isinstance(result, dict):
                st.session_state.summary = result.get("text", result.get("summary", ""))
            else:
                st.session_state.summary = result
            st.success("✅ Summary generated!")
        except Exception as e:
            st.error(f"❌ Summary generation failed: {e}")

if st.session_state.summary:
    with st.expander("📄 Summary"):
        st.write(st.session_state.summary)


if st.button("📚 Field of the Invention"):
    with st.spinner("Generating field of the invention..."):
        try:
            result = generate_field_of_invention(abstract)
            if isinstance(result, dict):
                st.session_state.field_of_invention = result.get("text", result.get("field", ""))
            else:
                st.session_state.field_of_invention = result
            st.success("Done!")
        except Exception as e:
            st.error(f"❌ Field generation failed: {e}")

if st.session_state.field_of_invention:
    with st.expander("📘 Field of the Invention"):
        st.write(st.session_state.field_of_invention)


if st.button("🧠 Background"):
    with st.spinner("Generating background..."):
        try:
            result = generate_background_locally(abstract)
            if isinstance(result, dict):
                st.session_state.background = result.get("text", result.get("background", ""))
            else:
                st.session_state.background = result
            st.success("Done!")
        except Exception as e:
            st.error(f"❌ Background generation failed: {e}")

if st.session_state.background:
    with st.expander("🔍 Background"):
        st.write(st.session_state.background)


if st.button("🎯 Objects of the Invention"):
    with st.spinner("Generating objects of the invention..."):
        try:
            result = generate_objects_of_invention(abstract)
            if isinstance(result, dict):
                text = result.get("text", result.get("objects", ""))
            else:
                text = result
            
            # Clean markdown headers and formatting
            import re
            text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)  # Remove markdown headers
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove bold
            text = re.sub(r'__([^_]+)__', r'\1', text)  # Remove underline
            
            st.session_state.objects_of_invention = text
            st.success("Done!")
        except Exception as e:
            st.error(f"❌ Objects generation failed: {e}")

if st.session_state.objects_of_invention:
    with st.expander("🎯 Objects of the Invention"):
        st.text(st.session_state.objects_of_invention)  # Use st.text instead of st.write


if st.button("📝 Detailed Description"):
    # Check if claims exist first
    if not st.session_state.get("claims"):
        st.warning("⚠️ Please generate Claims first!")
    else:
        with st.spinner("Generating detailed description (this may take 30-60 seconds)..."):
            try:
                # Get claims and drawing summary
                claims_text = st.session_state.get("claims", "Claims not generated yet.")
                drawings_text = drawing_summary if drawing_summary and drawing_summary.strip() else "No drawings provided."
                
                # Debug info
                st.write(f"**Debug Info:**")
                st.write(f"- Abstract length: {len(abstract)} chars")
                st.write(f"- Claims length: {len(claims_text)} chars")
                st.write(f"- Drawings length: {len(drawings_text)} chars")
                
                # Generate
                result = generate_detailed_description(
                    abstract,
                    claims_text,
                    drawings_text
                )
                
                # Handle result
                if isinstance(result, dict):
                    text = result.get("text", result.get("description", ""))
                else:
                    text = result
                
                # Debug output
                st.write(f"**Generated length:** {len(text)} characters")
                
                # Store
                if text and len(text) > 50:
                    st.session_state.detailed_description = text
                    st.success("✅ Done! Scroll down to see the detailed description.")
                else:
                    st.session_state.detailed_description = "⚠️ Generated description is too short or empty."
                    st.error("⚠️ Generated description is too short. Check model output.")
                    
            except Exception as e:
                error_msg = f"❌ Exception: {str(e)}"
                st.session_state.detailed_description = error_msg
                st.error(f"❌ Detailed description generation failed: {e}")
                st.exception(e)  # Show full traceback

if st.session_state.get("detailed_description") and len(st.session_state.get("detailed_description", "")) > 50:
    with st.expander("📘 Detailed Description", expanded=True):
        st.markdown(st.session_state.detailed_description)

if st.button("📊 Brief Description of Drawings"):
    if not abstract or not drawing_summary:
        st.warning("Please enter both the abstract and drawing summary.")
    else:
        with st.spinner("Generating brief description of drawings..."):
            try:
                result = generate_brief_description(abstract, drawing_summary)
                if isinstance(result, dict):
                    st.session_state.brief_description = result.get("text", result.get("description", ""))
                else:
                    st.session_state.brief_description = result or "⚠️ No output generated."
                st.success("Done!")
            except Exception as e:
                st.error(f"❌ Brief description generation failed: {e}")

if st.session_state.get("brief_description"):
    with st.expander("🖼️ Brief Description of the Drawings"):
        st.write(st.session_state.brief_description)


if st.button("🖼️ Summary of Drawings"):
    if not abstract:
        st.warning("Please enter the invention abstract.")
    else:
        with st.spinner("Generating summary of drawings..."):
            try:
                result = generate_drawing(abstract)
                if isinstance(result, dict):
                    st.session_state.summary_drawings = result.get("text", "")
                else:
                    st.session_state.summary_drawings = result or "⚠️ No output generated."
                st.success("Done!")
            except Exception as e:
                st.error(f"❌ Drawing summary failed: {e}")

if st.session_state.get("summary_drawings"):
    with st.expander("📷 Summary of Drawings"):
        st.write(st.session_state.summary_drawings)


# ------------------ CPC CLASSIFIER ----------------------
st.markdown("## 📚 CPC Classifier")
if st.button("🏷️ Classify CPC"):
    with st.spinner("Classifying CPC..."):
        try:
            result = classify_cpc(abstract)
            st.session_state.cpc_result = result or "⚠️ No result."
            st.success("Done!")
        except Exception as e:
            st.session_state.cpc_result = f"❌ Exception: {e}"
            st.error(f"❌ CPC classification failed: {e}")

if st.session_state.get("cpc_result"):
    with st.expander("🔍 CPC Classification"):
        st.code(st.session_state.cpc_result)

st.markdown("---")
st.markdown("## 🔍 Patent Quality Verification")
st.info("🤖 5 AI Agents will analyze your patent for quality and compliance")

if st.button("✅ Run 5-Agent Verification"):
    # Check if required sections exist
    required = ['title', 'claims', 'abstract_input', 'background', 'summary']
    missing = [s for s in required if not st.session_state.get(s)]
    
    if missing:
        st.warning(f"⚠️ Please generate these sections first: {', '.join(missing)}")
    else:
        with st.spinner("🤖 5 AI Agents verifying your patent... This may take 1-2 minutes"):
            
            try:
                from patent_verifier import verify_patent_5_sections
                
                # Prepare 5 critical sections
                sections_to_verify = {
                    'title': st.session_state.get('title', ''),
                    'abstract': st.session_state.get('abstract_input', ''),
                    'claims': st.session_state.get('claims', ''),
                    'background': st.session_state.get('background', ''),
                    'summary': st.session_state.get('summary', '')
                }
                
                # Run verification (this is where the real work happens)
                report = verify_patent_5_sections(sections_to_verify)
                
                # Display results
                st.success("✅ Verification Complete!")
                
                with st.expander("📊 Verification Report", expanded=True):
                    st.text(report)  # ✅ CHANGED: st.write() → st.text()
                
                # Store report
                st.session_state.verification_report = report
                
            except ImportError:
                st.error("❌ Patent verifier module not found")
                st.info("💡 Make sure patent_verifier.py is in the same directory")
            except Exception as e:
                st.error(f"❌ Verification failed: {str(e)}")
                st.info("💡 Check that Ollama is running: ollama serve")

# Show previous report if exists
if st.session_state.get('verification_report'):
    with st.expander("📋 View Previous Verification Report"):
        st.text(st.session_state.verification_report)  # ✅ CHANGED: st.write() → st.text()

st.markdown("---")

# ----------------------- EXPORT -------------------------
# ----------------------- EXPORT -------------------------
st.markdown("---")
st.markdown("## 📄 Export Patent Document")

export_abstract = st.session_state.get("abstract_input", "")

# ✅ CORRECT ORDER: Indian Patent Office Standard Structure
pdf_sections = {
    "Abstract": export_abstract or "[Not Provided]",  # 1. Abstract FIRST
    "Title": st.session_state.get("title", "[Not Generated]"),  # 2. Title
    "Field of the Invention": st.session_state.get("field_of_invention", "[Not Generated]"),  # 3. Field
    "Background of the Invention": st.session_state.get("background", "[Not Generated]"),  # 4. Background
    "Objects of the Invention": st.session_state.get("objects_of_invention", "[Not Generated]"),  # 5. Objects
    "Summary of the Invention": st.session_state.get("summary", "[Not Generated]"),  # 6. Summary ✅ ADDED
    "Brief Description of the Drawings": st.session_state.get("brief_description", "[Not Generated]"),  # 7. Brief Desc
    "Detailed Description of the Invention": st.session_state.get("detailed_description", "[Not Generated]"),  # 8. Detailed
    "Claims": st.session_state.get("claims", "[Not Generated]"),  # 9. Claims LAST
}

st.session_state.generated_sections = pdf_sections

# Show preview with section order
with st.expander("📋 Preview Export Sections (Indian Patent Office Order)"):
    section_list = [
        ("1️⃣", "Abstract"),
        ("2️⃣", "Title"),
        ("3️⃣", "Field of the Invention"),
        ("4️⃣", "Background of the Invention"),
        ("5️⃣", "Objects of the Invention"),
        ("6️⃣", "Summary of the Invention"),
        ("7️⃣", "Brief Description of the Drawings"),
        ("8️⃣", "Detailed Description of the Invention"),
        ("9️⃣", "Claims")
    ]
    
    for emoji, section_name in section_list:
        content = pdf_sections.get(section_name, "[Not Generated]")
        status = "✅" if content and content != "[Not Generated]" and content != "[Not Provided]" else "❌"
        word_count = len(content.split()) if content and content not in ["[Not Generated]", "[Not Provided]"] else 0
        st.write(f"{emoji} {status} **{section_name}**: {word_count} words")

col1, col2 = st.columns(2)

with col1:
    st.markdown("### 🧾 Generate PDF")
    if export_abstract.strip():
        if st.button("📄 Generate PDF"):
            with st.spinner("Creating PDF..."):
                try:
                    pdf_path = create_patent_pdf(pdf_sections)
                    st.success("✅ PDF Generated!")
                    with open(pdf_path, "rb") as f:
                        st.download_button(
                            label="📥 Download Patent PDF",
                            data=f,
                            file_name="patent_document.pdf",
                            mime="application/pdf"
                        )
                except Exception as e:
                    st.error(f"❌ PDF generation failed: {e}")
                    st.info("💡 Check that create_patent_pdf() function exists and is working")
    else:
        st.warning("⚠️ Please enter an abstract before generating the PDF.")

with col2:
    st.markdown("### 📝 Export as DOCX")
    if export_abstract.strip():
        if st.button("📝 Generate Indian Patent Office DOCX"):
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from io import BytesIO
                
                doc = Document()
                
                # Set margins (Indian Patent Office standard)
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(1.0)
                    section.bottom_margin = Inches(1.0)
                    section.left_margin = Inches(1.0)
                    section.right_margin = Inches(1.0)
                
                # Set default font
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                paragraph_format = style.paragraph_format
                paragraph_format.space_after = Pt(12)
                paragraph_format.line_spacing = 1.5
                
                # ========== PAGE 1: ABSTRACT (Standalone) ==========
                abstract_header = doc.add_paragraph()
                abstract_header_run = abstract_header.add_run("ABSTRACT")
                abstract_header_run.bold = True
                abstract_header_run.font.name = 'Times New Roman'
                abstract_header_run.font.size = Pt(12)
                
                abstract_content = export_abstract or "[Not Provided]"
                abstract_para = doc.add_paragraph(abstract_content.strip())
                abstract_para.style = doc.styles['Normal']
                
                doc.add_page_break()  # New page after abstract
                
                # ========== PAGE 2+: TITLE (Centered) ==========
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.add_run(st.session_state.get("title", "TITLE OF THE INVENTION"))
                title_run.bold = True
                title_run.font.name = 'Times New Roman'
                title_run.font.size = Pt(14)
                
                doc.add_paragraph()
                doc.add_paragraph()
                
                # ========== ALL OTHER SECTIONS IN ORDER ==========
                patent_sections = [
                    ("FIELD OF THE INVENTION", "field_of_invention"),
                    ("BACKGROUND OF THE INVENTION", "background"),
                    ("OBJECTS OF THE INVENTION", "objects_of_invention"),
                    ("SUMMARY OF THE INVENTION", "summary"),  # ✅ INCLUDED
                    ("BRIEF DESCRIPTION OF THE DRAWINGS", "brief_description"),
                    ("DETAILED DESCRIPTION OF THE INVENTION WITH REFERENCE TO THE ACCOMPANYING FIGURES", "detailed_description"),
                ]
                
                for section_title, session_key in patent_sections:
                    # Section header (Bold, uppercase)
                    header = doc.add_paragraph()
                    header_run = header.add_run(section_title)
                    header_run.bold = True
                    header_run.font.name = 'Times New Roman'
                    header_run.font.size = Pt(12)
                    
                    # Get content
                    content = st.session_state.get(session_key, "[Not Generated]")
                    
                    # Add content
                    content_para = doc.add_paragraph(content.strip() if content and content.strip() else "[Not Generated]")
                    content_para.style = doc.styles['Normal']
                    
                    doc.add_paragraph()  # Spacing between sections
                
                # ========== FINAL SECTION: CLAIMS (with WE CLAIM) ==========
                doc.add_page_break()  # Claims on new page (optional but professional)
                
                claims_header = doc.add_paragraph()
                claims_header_run = claims_header.add_run("CLAIMS")
                claims_header_run.bold = True
                claims_header_run.font.name = 'Times New Roman'
                claims_header_run.font.size = Pt(12)
                
                doc.add_paragraph()
                
                claims_content = st.session_state.get("claims", "[Not Generated]")
                if claims_content and claims_content != "[Not Generated]":
                    # "WE CLAIM" for Indian Patent Office
                    we_claim_para = doc.add_paragraph()
                    we_claim_run = we_claim_para.add_run("WE CLAIM")
                    we_claim_run.bold = True
                    we_claim_run.font.name = 'Times New Roman'
                    we_claim_run.font.size = Pt(12)
                    
                    doc.add_paragraph()
                    
                    claims_para = doc.add_paragraph(claims_content.strip())
                    claims_para.style = doc.styles['Normal']
                else:
                    not_gen_para = doc.add_paragraph("[Not Generated]")
                    not_gen_para.style = doc.styles['Normal']
                
                # Save to buffer
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 Download Indian Patent Office DOCX",
                    data=buffer,
                    file_name="patent_application_indian_format.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("✅ Indian Patent Office DOCX generated successfully!")
                st.info("📋 Structure: Abstract (Page 1) → Title → 7 Sections → Claims (Final)")
                
            except Exception as e:
                st.error(f"❌ DOCX generation failed: {e}")
                st.info("💡 Make sure python-docx is installed: pip install python-docx")
                import traceback
                st.code(traceback.format_exc())
    else:
        st.warning("⚠️ Please enter an abstract before generating the DOCX.")

st.markdown("---")

if st.button("🔄 Reset All"):
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.rerun()
